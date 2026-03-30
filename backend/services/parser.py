import pdfplumber
import docx
import io
from pathlib import Path

def parse_file_path(path: Path):
    ext = path.suffix.lower()
    if ext == ".pdf":
        import pdfplumber
        parts = []
        with pdfplumber.open(str(path)) as pdf:
            for p in pdf.pages:
                t = p.extract_text() or ""
                if t.strip():
                    parts.append(t)
        return "\n\n".join(parts), "pdf"

    if ext == ".docx":
        import docx
        d = docx.Document(str(path))
        return "\n".join([p.text for p in d.paragraphs if p.text.strip()]), "docx"

    if ext == ".txt":
        return path.read_text(encoding="utf-8", errors="ignore"), "txt"

    return "", "unknown"


async def parse_file(file):
    filename = file.filename
    file_type = filename.split(".")[-1].lower() if filename else "unknown"
    content = ""
    
    # Read file content once
    file_content = await file.read()
    
    if not file_content:
        return "", file_type

    try:
        if file_type == "pdf":
            with pdfplumber.open(io.BytesIO(file_content)) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        content += text + "\n"
        elif file_type == "docx":
            doc = docx.Document(io.BytesIO(file_content))
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            content = "\n".join(paragraphs)
        elif file_type == "txt":
            try:
                content = file_content.decode("utf-8")
            except UnicodeDecodeError:
                # Try with error handling
                content = file_content.decode("utf-8", errors="ignore")
        else:
            # Try to decode as text for unknown types
            try:
                content = file_content.decode("utf-8", errors="ignore")
            except:
                content = ""
    except Exception as e:
        print(f"Error parsing file {filename}: {e}")
        return "", file_type

    return content.strip(), file_type
